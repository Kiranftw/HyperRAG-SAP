import os
import sys
# Allow importing RAG when running directly from the agents/ folder
parent_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if parent_dir not in sys.path:
    sys.path.append(parent_dir)
rag_dir = os.path.join(parent_dir, "rag")
if rag_dir not in sys.path:
    sys.path.append(rag_dir)
import json
import csv
import subprocess
import docx
import shlex
import pandas as pd
import cohere
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
from typing import Dict, List, TypedDict, Literal, Optional, Any, Union, Sequence
from pydantic import BaseModel, Field, field_validator
from rag.agentic_rag import AgenticRAG, HybridSearch, HyperRetrivalAugmentedGeneration, FAISSIndexGeneration, LOGGER, ExceptionHandelling
from langchain_community.document_loaders import (
    CSVLoader,
    JSONLoader,
    PyPDFLoader,
    TextLoader,
    Docx2txtLoader,
    DirectoryLoader,
    UnstructuredWordDocumentLoader,
    UnstructuredMarkdownLoader,
    UnstructuredHTMLLoader
)
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_core.output_parsers import SimpleJsonOutputParser
from langchain_huggingface import HuggingFaceEmbeddings
from elasticsearch import Elasticsearch, helpers
from langchain_ollama import ChatOllama
import requests
from bs4 import BeautifulSoup
from dotenv import find_dotenv, load_dotenv
from geopy.geocoders import Nominatim
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_nvidia_ai_endpoints import ChatNVIDIA
from langgraph.graph import END, StateGraph
from langchain_tavily import TavilySearch
from langchain_community.document_loaders import PyMuPDFLoader
load_dotenv(find_dotenv())

ALLOWED_FILE_TYPES = {
    ".pdf",
    ".docx",
    ".txt",
    ".json",
    ".png",
    ".jpg",
    ".jpeg",
}
MAX_FILE_SIZE_MB = 50

class DocumentIngestionRequest(BaseModel):
    filename: str = Field(..., description="Original uploaded filename")
    filepath: str = Field(..., description="Temporary or permanent file path")
    filesize_mb: float = Field(..., gt=0)
    content_type: str = Field(..., description="MIME type")
    uploaded_by: Optional[str] = None
    enable_ocr: bool = False
    enable_chunking: bool = True
    enable_embedding: bool = True
    chunk_size: int = Field(default=1000, ge=100, le=4000)
    chunk_overlap: int = Field(default=200, ge=0, le=1000)
    metadata: Optional[Dict] = {}
    @field_validator("filename")
    @classmethod
    def validate_extension(cls, value: str):
        ext = os.path.splitext(value)[1].lower()
        if ext not in ALLOWED_FILE_TYPES:
            raise ValueError(
                f"Unsupported file type: {ext}"
            )
        return value
    @field_validator("filesize_mb")
    @classmethod
    def validate_file_size(cls, value: float):
        if value > MAX_FILE_SIZE_MB:
            raise ValueError(
                f"File exceeds max limit of {MAX_FILE_SIZE_MB} MB"
            )
        return value

class SearchInternet(BaseModel):
    query: str = Field(..., description="search query")
    documents_count: int = Field(..., description="number of document to retrieve")

class QueryDecomposition(BaseModel):
    query: str = Field(..., description="search query")
    queries_count: int = Field(..., description="Number of queries to decompose")

class SaveDocumentRequest(BaseModel):
    filename: str = Field(..., description="The name of the file to save, e.g., 'report.json' or 'guide.txt'")
    data: Any = Field(..., description="The content/data to save (raw string, list, or JSON dictionary)")
    font_family: str = Field(default="Arial, sans-serif", description="Optional font family for PDF output")
    font_size: str = Field(default="14px", description="Optional font size for PDF output")

class TerminalCommandRequest(BaseModel):
    command: str = Field(..., description="The shell command to execute, e.g. 'docker-compose up -d'")
    timeout: int = Field(default=120, description="Max seconds to wait before killing the command")

MODEL = ChatNVIDIA(
    model="moonshotai/kimi-k2-instruct",
    api_key=os.getenv("NVIDIA_API_KEY"),
    temperature=1,
    top_p=0.9,
    max_completion_tokens=16384,
)
ALLOWED_COMMANDS = {"python", "python3", "pytest", "git", "npm", "node", "pip", "powershell", "cmd", "bash", "sh", "uv", "pnpm", "ls", "cd", "pwd", "mkdir", "rm", "touch", "echo", "curl", "docker", "docker-compose", "kubectl", "minikube"}
FORBIDDEN_TOKENS = {"&", "|", ";", "`", "$(", ">", "<"}

def _normalize_command(command: Union[str, Sequence[str]]) -> List[str]:
    if isinstance(command, str):
        if any(tok in command for tok in FORBIDDEN_TOKENS):
            raise ValueError("Shell operators are not allowed.")
        arguments = shlex.split(command, posix=(os.name != "nt"))
    else:
        arguments = list(command)
    if not arguments:
        raise ValueError("Empty command.")
    exe = os.path.basename(arguments[0]).lower()
    if exe.endswith(".exe"):
        exe = exe[:-4]
    if exe not in ALLOWED_COMMANDS:
        raise ValueError(f"Command '{arguments[0]}' is not allowed.")
    return arguments

class AgentTools(AgenticRAG):
    def __init__(self):
        super().__init__()

    def search_internet(self, request: SearchInternet) -> Dict:
        """Search the internet for SAP documentation or general information using Tavily."""
        query = request.query
        docunments_count = request.documents_count
        SEARCH_ENGINE = TavilySearch(
            tavily_api_key=os.getenv("TAVILY_API_KEY"),
            max_results=docunments_count,
            include_answer=True,
            include_raw_content=True,
            include_urls=True,
            include_tables=True,
            include_domains=[
                "help.sap.com",
                "www.sap.com",
                "developers.sap.com",
                "api.sap.com",
                "community.sap.com",
            ],
            include_images=True,
        )
        try:
            response = SEARCH_ENGINE.invoke(query)
            if isinstance(response, str):
                try:
                    response = json.loads(response)
                except Exception:
                    response = {"results": [], "answer": response, "images": []}
            if not isinstance(response, dict):
                response = {"results": [], "answer": str(response), "images": []}

            normalized_results = []
            for item in response.get("results", []):
                normalized_results.append({
                    "title": item.get("title"),
                    "url": item.get("url"),
                    "content": item.get("content"),
                })
            return {
                "query": query,
                "summary": response.get("answer", ""),
                "images": response.get("images", []),
                "sources": normalized_results,
            }
        except Exception as e:
            return {
                "error": str(e)
            }

    def query_decomposition(self, request: QueryDecomposition) -> List[str]:
        """Decompose a complex query into simpler sub-queries for better retrieval."""
        # Load decomposition prompt
        with open(
            os.path.join(self.ROOT, "prompts", "query_decomposition.txt"), "r"
        ) as file:
            decomposition_prompt = file.read()
        # Fill the prompt with the user query
        decomposition_prompt = decomposition_prompt.replace(
            "{USER_QUERY_GOES_HERE}", request.query
        )
        try:
            DECOMPOSITION_QUERIES_COUNT: int = request.queries_count
            client = getattr(self, "cohere_reranker", None) or cohere.ClientV2(
                os.getenv("COHERE_API_KEY")
            )
            response = client.chat(
                model="command-r-plus",
                messages=[
                    {
                        "role": "system",
                        "content": decomposition_prompt
                        + f"The number of queries should be {DECOMPOSITION_QUERIES_COUNT}",
                    },
                    {"role": "user", "content": request.query},
                ],
                response_format={"type": "json_object"},
                temperature=0.7,
            )
            # Parse the JSON response
            content = response.message.content[0].text
            # Sometimes LLMs wrap JSON in markdown blocks
            if content.startswith("```json"):
                content = content.replace("```json\n", "").replace("```", "").strip()
            elif content.startswith("```"):
                content = content.replace("```\n", "").replace("```", "").strip()
            parsed_json = json.loads(content)
            sub_queries = parsed_json.get("sub_queries", [request.query])
            return sub_queries
        except Exception as e:
            print(f"Error during query decomposition: {e}")
            return [request.query]

    def write_file(self, request: SaveDocumentRequest) -> Dict:
        """Save extracted data, lists, or text to a local file (supports JSON, PDF, DOCX, XLSX, CSV, TXT)."""
        try:
            folder = "datasets"
            if not os.path.exists(folder):
                os.makedirs(folder)
            filepath = os.path.join(folder, request.filename)
            # Determine format based on requested filename extension
            ext = os.path.splitext(request.filename)[1].lower()
            # If no extension is specified, guess based on data type
            if not ext:
                if isinstance(request.data, (dict, list)):
                    ext = ".json"
                    filepath += ".json"
                else:
                    ext = ".txt"
                    filepath += ".txt"

            if ext == ".json":
                with open(filepath, "w", encoding="utf-8") as f:
                    # Support saving both serializable structures and JSON strings
                    if isinstance(request.data, (dict, list)):
                        json.dump(request.data, f, indent=4)
                    else:
                        try:
                            parsed = json.loads(str(request.data))
                            json.dump(parsed, f, indent=4)
                        except Exception:
                            # Write raw string if parse fails
                            f.write(str(request.data))

            elif ext == ".pdf":
                import markdown
                from weasyprint import HTML
                text_content = str(request.data)
                font_family = getattr(request, 'font_family', 'Arial, sans-serif')
                font_size = getattr(request, 'font_size', '14px')
                # Convert Markdown to HTML
                html_content = markdown.markdown(text_content, extensions=['tables', 'fenced_code'])
                # Basic styling for the PDF
                css_style = f"""
                    body {{ font-family: {font_family}; font-size: {font_size}; padding: 20px; line-height: 1.6; color: #333; }}
                    h1, h2, h3 {{ color: #0056b3; }}
                    table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; }}
                    code {{ background-color: #f8f9fa; padding: 2px 4px; border-radius: 4px; }}
                    pre {{ background-color: #f8f9fa; padding: 15px; border-radius: 4px; overflow-x: auto; }}
                """
                full_html = f"<html><head><style>{css_style}</style></head><body>{html_content}</body></html>"
                # Generate PDF
                HTML(string=full_html).write_pdf(filepath)
            elif ext == ".docx":
                doc = docx.Document()
                text_content = str(request.data)
                paragraphs = text_content.split("\n")
                for p_text in paragraphs:
                    if p_text.strip():
                        doc.add_paragraph(p_text)
                    else:
                        doc.add_paragraph("")
                doc.save(filepath)

            elif ext in [".xlsx", ".xls"]:
                data = request.data
                if isinstance(data, list):
                    if len(data) > 0 and isinstance(data[0], dict):
                        df = pd.DataFrame(data)
                    elif len(data) > 0 and isinstance(data[0], list):
                        df = pd.DataFrame(data)
                    else:
                        df = pd.DataFrame({"Data": data})
                elif isinstance(data, dict):
                    df = pd.DataFrame([data])
                else:
                    try:
                        parsed = json.loads(str(data))
                        if isinstance(parsed, list):
                            df = pd.DataFrame(parsed)
                        elif isinstance(parsed, dict):
                            df = pd.DataFrame([parsed])
                        else:
                            df = pd.DataFrame({"Content": [str(data)]})
                    except Exception:
                        df = pd.DataFrame({"Content": [str(data)]})
                df.to_excel(filepath, index=False)
            elif ext == ".csv":
                # Save as CSV file
                data = request.data
                if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
                    keys = data[0].keys()
                    with open(filepath, "w", newline="", encoding="utf-8") as f:
                        writer = csv.DictWriter(f, fieldnames=keys)
                        writer.writeheader()
                        writer.writerows(data)
                elif isinstance(data, list) and len(data) > 0 and isinstance(data[0], list):
                    with open(filepath, "w", newline="", encoding="utf-8") as f:
                        writer = csv.writer(f)
                        writer.writerows(data)
                else:
                    # Raw string or fallback format
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(str(data))
            else:
                # Default fallback (txt, md, log, etc.)
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(str(request.data))
            LOGGER.info(f"Document saved successfully to {filepath}")
            return {
                "status": "success",
                "filename": os.path.basename(filepath),
                "filepath": filepath,
                "message": f"Document saved successfully as {ext.upper()}"
            }
        except Exception as e:
            LOGGER.error(f"Error saving document: {e}")
            return {
                "status": "error",
                "message": f"Error saving document: {str(e)}"
            }

    @staticmethod
    def process_urls(urls: List[str]) -> List[Dict]:
        """Scrape and extract clean text content from a list of web URLs."""
        tavily_api_key = os.getenv("TAVILY_API_KEY")
        if tavily_api_key:
            try:
                LOGGER.info("Attempting URL content extraction using Tavily Extract API...")
                url = "https://api.tavily.com/extract"
                TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")
                if not TAVILY_API_KEY:
                    LOGGER.warning("TAVILY_API_KEY not found in environment variables. Skipping Tavily Extract API.")
                    return []
                payload = {
                    "urls": urls,
                    "extract_depth": "advanced"
                }
                headers = {
                    "Authorization": f"Bearer {TAVILY_API_KEY}",
                    "Content-Type": "application/json"
                }
                response = requests.post(url, json=payload, headers=headers, timeout=20)
                if response.status_code == 200:
                    data = response.json()
                    results = []
                    # Parse successfully extracted results
                    for item in data.get("results", []):
                        results.append({
                            "status": "success",
                            "url": item.get("url"),
                            "title": item.get("title") or "No Title",
                            "extracted_content": item.get("raw_content") or ""
                        })
                    # Track failed results to retry via local fallback
                    failed_urls = [item.get("url") for item in data.get("failed_results", []) if item.get("url")]
                    # If everything succeeded, return results
                    if results and not failed_urls:
                        return results
                    # Otherwise, fallback for failed/missing URLs
                    extracted_urls = {r["url"] for r in results}
                    urls_to_fallback = [u for u in urls if u not in extracted_urls]
                    if urls_to_fallback:
                        LOGGER.warning(f"Tavily Extract failed or skipped {len(urls_to_fallback)} URLs. Using local scraper fallback.")
                        fallback_results = self._fallback_scraping(urls_to_fallback)
                        results.extend(fallback_results)
                    return results
                else:
                    LOGGER.warning(f"Tavily Extract API returned status code {response.status_code}: {response.text}")
            except Exception as e:
                LOGGER.warning(f"Error during Tavily Extract API call: {e}")

        LOGGER.info("Using local scraping fallback...")
        return self._fallback_scraping(urls)

    def _fallback_scraping(self, urls: List[str]) -> List[Dict]:
        results = []
        for url in urls:
            # 1. Try newspaper3k first as it uses heuristic parsing and custom user-agents
            try:
                from newspaper import Article
                LOGGER.info(f"Trying newspaper3k parsing for URL: {url}")
                article = Article(url)
                article.download()
                article.parse()
                title = article.title
                text = article.text
                if text and len(text.strip()) > 100 and "Access Denied" not in title:
                    results.append({
                        "status": "success",
                        "url": url,
                        "title": title,
                        "extracted_content": text[:8000]
                    })
                    continue
            except Exception as e:
                LOGGER.debug(f"newspaper3k failed for {url}: {e}. Falling back to BeautifulSoup.")
            # 2. Fall back to standard BeautifulSoup scraping with robust modern browser headers
            try:
                LOGGER.info(f"Locally scraping URL with BeautifulSoup: {url}")
                headers = {
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8",
                    "Accept-Language": "en-US,en;q=0.9",
                    "Accept-Encoding": "gzip, deflate, br",
                    "Connection": "keep-alive",
                    "Upgrade-Insecure-Requests": "1",
                    "Sec-Ch-Ua": '"Not_A Brand";v="8", "Chromium";v="120", "Google Chrome";v="120"',
                    "Sec-Ch-Ua-Mobile": "?0",
                    "Sec-Ch-Ua-Platform": '"Windows"',
                    "Sec-Fetch-Dest": "document",
                    "Sec-Fetch-Mode": "navigate",
                    "Sec-Fetch-Site": "none",
                    "Sec-Fetch-User": "?1"
                }
                response = requests.get(url, headers=headers, timeout=15)
                if response.status_code in [401, 403]:
                    results.append({
                        "status": "error",
                        "url": url,
                        "error": f"Access Denied (HTTP {response.status_code}). WAF/Akamai/Cloudflare block detected."
                    })
                    continue
                soup = BeautifulSoup(response.text, "html.parser")
                title = soup.title.string.strip() if soup.title and soup.title.string else "No Title"
                if "Access Denied" in title or "Attention Required" in title or "Cloudflare" in title:
                    results.append({
                        "status": "error",
                        "url": url,
                        "error": "WAF / Cloudflare block page detected in title."
                    })
                    continue
                for script in soup(["script", "style", "nav", "footer", "header", "noscript", "aside"]):
                    script.decompose()
                clean_text = soup.get_text(separator="\n").strip()
                clean_text = "\n".join([line.strip() for line in clean_text.splitlines() if line.strip()])
                results.append({
                    "status": "success",
                    "url": url,
                    "title": title,
                    "extracted_content": clean_text[:8000]
                })
            except Exception as e:
                LOGGER.error(f"Error locally scraping {url}: {e}")
                results.append({
                    "status": "error",
                    "url": url,
                    "error": str(e)
                })
        return results

    # Reading Files from the Local System
    @ExceptionHandelling
    def read_file(self, filepath: str) -> str:
        """Read text contents of a local file (Supports PDF, TXT, DOCX, and Image OCR)."""
        if not filepath:
            LOGGER.error("NO FILEPATH IS AVAILABLE!")
            return ""
        if not os.path.exists(filepath):
            LOGGER.error(f"FILE NOT FOUND: {filepath}")
            return ""
        
        extracted_data = ""
        filename: str = os.path.basename(filepath)

        if filename.lower().endswith('.pdf'):
            loader: PyMuPDFLoader = PyMuPDFLoader(filepath)
            documents = loader.load()
            extracted_data = '\n'.join(doc.page_content for doc in documents)
            # If no PDF), use OCR with PyMuPDF
            if len(extracted_data.strip()) < 50:
                LOGGER.info("NO TEXT EXTRAC-TED FROM PDF, USING OCR...(SCANNED PDF)")
                doc = fitz.open(filepath)
                ocr_texts = []
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    # Convert page to image at 300 DPI
                    pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
                    img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                    # OCR the image
                    text = pytesseract.image_to_string(img, lang='eng')
                    ocr_texts.append(text)
                    LOGGER.info(f"OCR page {page_num + 1}: EXTRACTED {len(text)} CHARACTERS")
                doc.close()
                extracted_data = '\n\n'.join(ocr_texts)
                LOGGER.info(f"TOTAL OCR CHARACTERS: {len(extracted_data)} CHARACTERS")
            return extracted_data
    
        elif filename.lower().endswith('.txt'):
            loader: TextLoader = TextLoader(filepath)
            documents = loader.load()
            extracted_data = '\n'.join(doc.page_content for doc in documents)
            return extracted_data

        elif filename.lower().endswith('.docx'):
            loader: Docx2txtLoader = Docx2txtLoader(filepath)
            documents = loader.load()
            extracted_data = '\n'.join(doc.page_content for doc in documents)
            return extracted_data
        
        elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif')):
            image = Image.open(filepath)
            extracted_data = pytesseract.image_to_string(image)
            extracted_data = extracted_data.strip()
            LOGGER.info(f"IMAGE OCR: EXTRACTED {len(extracted_data)} CHARACTERS")
            return extracted_data
        else:
            LOGGER.error("UNSUPPORTED FILE TYPE!")
            return ""
    
    def bash(self, request: TerminalCommandRequest) -> Dict:
        """Execute an allowed terminal shell command safely and return stdout/stderr."""
        try:
            #subprocess run commands in detached mode without stdin, stdout, stderr attached to the subprocess (headless)
            args = _normalize_command(request.command)
            result = subprocess.run(
                args,
                shell=False,
                capture_output=True,
                text=True,
                timeout=request.timeout,
                cwd=getattr(self, "workspace_dir", None),
                # setting the env to make sure that we can run the command in the container
                # NOTE this env is for local development only
                env={
                    "PATH": os.environ.get("PATH", ""),
                    "HOME": os.environ.get("HOME", ""),
                    "USERPROFILE": os.environ.get("USERPROFILE", ""),
                },
            )
            return {
                "stdout": result.stdout,
                "stderr": result.stderr,
                "exit_code": result.returncode,
            }
        except subprocess.TimeoutExpired:
            return {
                "stdout": "",
                "stderr": f"Command timed out after {request.timeout}s",
                "exit_code": -1,
            }
        except Exception as e:
            LOGGER.error(f"Failed to run command: {e}")
            return {
                "stdout": "",
                "stderr": f"Error running command: {e}",
                "exit_code": -1,
            }