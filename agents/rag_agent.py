import os
import json
import csv
import docx
import pandas as pd
import cohere
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
from typing import Dict, List, TypedDict, Literal, Optional, Any
from pydantic import BaseModel, Field, field_validator
from RAG.AgenticRAG import AgenticRAG, HyperRetrivalAugmentedGeneration, FAISSIndexGeneration, LOGGER
import requests
from dotenv import find_dotenv, load_dotenv
from geopy.geocoders import Nominatim
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_nvidia_ai_endpoints import ChatNVIDIA
from langgraph.graph import END, StateGraph
from langchain_tavily import TavilySearch
from langchain_community.document_loaders import PyMuPDFLoader, TextLoader, Docx2txtLoader

load_dotenv(find_dotenv())

ALLOWED_FILE_TYPES = {
    ".pdf",
    ".docx",
    ".txt",
    ".json",
    ".png",
    ".jpg",
    ".jpeg",
}

MAX_FILE_SIZE_MB = 50

class DocumentIngestionRequest(BaseModel):
    filename: str = Field(..., description="Original uploaded filename")
    filepath: str = Field(..., description="Temporary or permanent file path")
    filesize_mb: float = Field(..., gt=0)
    content_type: str = Field(..., description="MIME type")
    uploaded_by: Optional[str] = None
    enable_ocr: bool = False
    enable_chunking: bool = True
    enable_embedding: bool = True

    chunk_size: int = Field(default=1000, ge=100, le=4000)
    chunk_overlap: int = Field(default=200, ge=0, le=1000)
    metadata: Optional[Dict] = {}

    @field_validator("filename")
    @classmethod
    def validate_extension(cls, value: str):
        ext = os.path.splitext(value)[1].lower()
        if ext not in ALLOWED_FILE_TYPES:
            raise ValueError(
                f"Unsupported file type: {ext}"
            )
        return value
    @field_validator("filesize_mb")
    @classmethod
    def validate_file_size(cls, value: float):
        if value > MAX_FILE_SIZE_MB:
            raise ValueError(
                f"File exceeds max limit of {MAX_FILE_SIZE_MB} MB"
            )
        return value

class SearchInternet(BaseModel):
    query: str = Field(..., description="search query")
    @field_validator("query")
    @classmethod
    def validate_query(cls, value: str):
        if not value:
            raise ValueError("Query cannot be empty")
        return value
    documents_count: int = Field(..., description="number of document to retrieve")
    @field_validator("documents_count")
    @classmethod
    def validate_documents_count(cls, value: int):
        if value <= 0:
            raise ValueError("number of documents must be > 0")
        return value
    
class QueryDecomposition(BaseModel):
    query: str = Field(..., description="search query")
    queries_count: int = Field(..., description="Number of queries to decompose")
    @field_validator("query")
    @classmethod
    def validate_query(cls, value: str):
        if not value:
            raise ValueError("Query cannot be empty")
        return value
    @field_validator("queries_count")
    @classmethod
    def validate_queries_count(cls, value: int):
        if value <= 0:
            raise ValueError("Number of queries must be greater than 0")
        return value

class SaveDocumentRequest(BaseModel):
    filename: str = Field(..., description="The name of the file to save, e.g., 'report.json' or 'guide.txt'")
    data: Any = Field(..., description="The content/data to save (raw string, list, or JSON dictionary)")

MODEL = ChatNVIDIA(
    model="moonshotai/kimi-k2-instruct",
    api_key=os.getenv("NVIDIA_API_KEY"),
    temperature=1,
    top_p=0.9,
    max_completion_tokens=16384,
)

class Tools(AgenticRAG):
    def __init__(self):
        super().__init__()
        
    def search_internet(self, request: SearchInternet) -> Dict:
        """
        Search the web for real-time information on a given topic.
        Use this tool when:
        - up-to-date information is needed
        - the answer requires internet search
        - the user asks about recent events or external knowledge
        Args:
            request (SearchInternet): The Pydantic request object.
        Returns:
            Dict containing search results with relevant web information.
        """
        query = request.query
        docunments_count = request.documents_count
        SEARCH_ENGINE = TavilySearch(
            tavily_api_key=os.getenv("TAVILY_API_KEY"),
            max_results=docunments_count,
            include_answer=True,
            include_raw_content=True,
            include_urls=True,
            include_tables=True,
            include_domains=[
                "help.sap.com",
                "www.sap.com",
                "developers.sap.com",
                "api.sap.com",
                "community.sap.com",
            ],
            include_images=True,
        )
        try:
            response: Dict = SEARCH_ENGINE.invoke(query)
            normalized_results = []
            for item in response.get("results", []):
                normalized_results.append({
                    "title": item.get("title"),
                    "url": item.get("url"),
                    "content": item.get("content"),
                })
            return {
                "query": query,
                "summary": response.get("answer", ""),
                "images": response.get("images", []),
                "sources": normalized_results,
            }
        except Exception as e:
            return {
                "error": str(e)
            }
    
    def query_decomposition(self, request: QueryDecomposition) -> List[str]:
        """
        Decompose a complex, multi-intent user query into multiple independent sub-queries.
        This method uses a Generative LLM to analyze the input query and break it down into
        atomic, modular questions. It enforces specific constraints such as expanding vague
        acronyms or pronouns into full domain names (e.g., "SAP S/4HANA Cloud Public Edition")
        to maximize exact-match keyword hits in the sparse BM25 retrieval stage.
        Args:
            request (QueryDecomposition): The request containing query and count.
        Returns:
            List[str]: A list of optimized, contextually independent sub-queries.
        """
        # Load decomposition prompt
        with open(
            os.path.join(self.ROOT, "prompts", "query_decomposition.txt"), "r"
        ) as file:
            decomposition_prompt = file.read()
        # Fill the prompt with the user query
        decomposition_prompt = decomposition_prompt.replace(
            "{USER_QUERY_GOES_HERE}", request.query
        )
        try:
            DECOMPOSITION_QUERIES_COUNT: int = request.queries_count
            client = getattr(self, "cohere_reranker", None) or cohere.ClientV2(
                os.getenv("COHERE_API_KEY")
            )
            response = client.chat(
                model="command-r-plus",
                messages=[
                    {
                        "role": "system",
                        "content": decomposition_prompt
                        + f"The number of queries should be {DECOMPOSITION_QUERIES_COUNT}",
                    },
                    {"role": "user", "content": request.query},
                ],
                response_format={"type": "json_object"},
                temperature=0.7,
            )
            # Parse the JSON response
            content = response.message.content[0].text
            # Sometimes LLMs wrap JSON in markdown blocks
            if content.startswith("```json"):
                content = content.replace("```json\n", "").replace("```", "").strip()
            elif content.startswith("```"):
                content = content.replace("```\n", "").replace("```", "").strip()
            parsed_json = json.loads(content)
            sub_queries = parsed_json.get("sub_queries", [request.query])
            return sub_queries
        except Exception as e:
            print(f"Error during query decomposition: {e}")
            return [request.query]
    
    def save_documents(self, request: SaveDocumentRequest) -> Dict:
        """
        Save text content, JSON data, PDF, Word, CSV, or Excel files into a file with a given filename under the 'uploaded_documents' directory.
        Use this tool when you need to persist search results, reports, guides, spreadsheets, or any generated data to a file.
        Args:
            request (SaveDocumentRequest): The request containing the filename and content data.
        Returns:
            Dict containing the file path of the saved document or error status.
        """
        try:
            folder = "datasets"
            if not os.path.exists(folder):
                os.makedirs(folder)
            filepath = os.path.join(folder, request.filename)
            # Determine format based on requested filename extension
            ext = os.path.splitext(request.filename)[1].lower()
            # If no extension is specified, guess based on data type
            if not ext:
                if isinstance(request.data, (dict, list)):
                    ext = ".json"
                    filepath += ".json"
                else:
                    ext = ".txt"
                    filepath += ".txt"
            
            if ext == ".json":
                with open(filepath, "w", encoding="utf-8") as f:
                    # Support saving both serializable structures and JSON strings
                    if isinstance(request.data, (dict, list)):
                        json.dump(request.data, f, indent=4)
                    else:
                        try:
                            parsed = json.loads(str(request.data))
                            json.dump(parsed, f, indent=4)
                        except Exception:
                            # Write raw string if parse fails
                            f.write(str(request.data))
                            
            elif ext == ".pdf":
                doc = fitz.open()
                text_content = str(request.data)
                lines = text_content.split("\n")
                page = doc.new_page()
                margin_left = 50
                margin_top = 50
                line_height = 15
                page_height = page.rect.height
                
                y = margin_top
                for line in lines:
                    if y + line_height > page_height - 50:
                        page = doc.new_page()
                        y = margin_top
                    page.insert_text((margin_left, y), line, fontsize=10)
                    y += line_height
                    
                doc.save(filepath)
                doc.close()
            elif ext == ".docx":
                doc = docx.Document()
                text_content = str(request.data)
                paragraphs = text_content.split("\n")
                for p_text in paragraphs:
                    if p_text.strip():
                        doc.add_paragraph(p_text)
                    else:
                        doc.add_paragraph("")
                doc.save(filepath)
                
            elif ext in [".xlsx", ".xls"]:
                data = request.data
                if isinstance(data, list):
                    if len(data) > 0 and isinstance(data[0], dict):
                        df = pd.DataFrame(data)
                    elif len(data) > 0 and isinstance(data[0], list):
                        df = pd.DataFrame(data)
                    else:
                        df = pd.DataFrame({"Data": data})
                elif isinstance(data, dict):
                    df = pd.DataFrame([data])
                else:
                    try:
                        parsed = json.loads(str(data))
                        if isinstance(parsed, list):
                            df = pd.DataFrame(parsed)
                        elif isinstance(parsed, dict):
                            df = pd.DataFrame([parsed])
                        else:
                            df = pd.DataFrame({"Content": [str(data)]})
                    except Exception:
                        df = pd.DataFrame({"Content": [str(data)]})
                
                df.to_excel(filepath, index=False)
            elif ext == ".csv":
                # Save as CSV file
                data = request.data
                if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
                    keys = data[0].keys()
                    with open(filepath, "w", newline="", encoding="utf-8") as f:
                        writer = csv.DictWriter(f, fieldnames=keys)
                        writer.writeheader()
                        writer.writerows(data)
                elif isinstance(data, list) and len(data) > 0 and isinstance(data[0], list):
                    with open(filepath, "w", newline="", encoding="utf-8") as f:
                        writer = csv.writer(f)
                        writer.writerows(data)
                else:
                    # Raw string or fallback format
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(str(data))           
            else:
                # Default fallback (txt, md, log, etc.)
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(str(request.data))
            LOGGER.info(f"Document saved successfully to {filepath}")
            return {
                "status": "success",
                "filename": os.path.basename(filepath),
                "filepath": filepath,
                "message": f"Document saved successfully as {ext.upper()}"
            }
        except Exception as e:
            LOGGER.error(f"Error saving document: {e}")
            return {
                "status": "error",
                "message": f"Error saving document: {str(e)}"
            }

TOOLS = Tools()

class AgentState(TypedDict):
    messages: List[Dict]
    user_input: str
    goal: str
    plan: List[str]
    current_action: str
    next_action: str
    tool_result: Dict
    tools_used: List[Dict]
    working_memory: Dict
    long_term_memory: Dict
    observations: List[Dict]
    last_error: str
    retry_count: int
    confidence: float
    status: str
    final_response: str
    source: Dict
    document_ids: List[str]
    ingestion_job_id: str
